"""
خدمة تصدير الردود إلى مستندات Word.
"""
import json
from pathlib import Path
from docx import Document
from docx.shared import Pt
from app.models.response import Response

EXPORT_DIR = Path("exports")
EXPORT_DIR.mkdir(exist_ok=True)

class ExportService:
    @staticmethod
    def export(response: Response, format: str = "docx") -> Path:
        if format != "docx":
            raise ValueError("تنسيق التصدير غير مدعوم حالياً")
        return ExportService._export_to_docx(response)

    @staticmethod
    def _export_to_docx(response: Response) -> Path:
        content = response.generated_content
        answers = []
        if isinstance(content, list):
            answers = content
        elif isinstance(content, str):
            try:
                answers = json.loads(content)
            except (json.JSONDecodeError, TypeError):
                answers = [{"question": "الرد الكامل", "answer": content}]

        doc = Document()
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Tajawal'
        font.size = Pt(12)

        doc.add_heading('رد العطاء', level=1)

        for item in answers:
            question = item.get('question', '')
            answer = item.get('answer', '')
            if question:
                doc.add_heading(question, level=2)
            doc.add_paragraph(answer)
            doc.add_paragraph()

        file_path = EXPORT_DIR / f"response_{response.analysis_id}.docx"
        doc.save(file_path)
        return file_path